import io
import os
import uuid
import calendar as cal_module
from datetime import datetime, timedelta, timezone

from dotenv import load_dotenv
load_dotenv()

from flask import (
    Flask, render_template, request, redirect,
    url_for, session, jsonify, Response, send_file, after_this_request
)

from docx import Document
from docx.shared import Pt
from werkzeug.security import generate_password_hash, check_password_hash

from database import init_db, get_db
import hemingway_client
import degas_client
import glossary
import postiz_client
import nextcloud_client

# ── Config ────────────────────────────────────────────────────────────────────
app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "change-me-in-production")

# Access codes (Ben's ask, 2026-09-03): replaced the single shared
# APP_PASSWORD with a table of labeled, independently-revocable passwords --
# "not necessarily user-based," just distinct secrets Ben can hand out per
# person/role and pull individually. See database.py's access_codes table
# for the schema/seeding, and manage_access_codes() below for the admin UI.
# This is NOT the full unified-auth story (STUDIO_SYSTEM_DESIGN.md Section
# 2/9, step 2) -- that requires Degas and Hemingway to accept Studio's
# session too, which is a separate, larger project. This is scoped to
# Studio only per Ben's explicit choice.


@app.before_request
def require_login():
    public = {"login", "static", "health"}
    if request.endpoint in public:
        return
    if not session.get("logged_in"):
        return redirect(url_for("login"))
    # Real-time revocation: checked on EVERY request, not just at login, so
    # revoking a code kicks out any session already using it immediately.
    code_id = session.get("access_code_id")
    if code_id is not None:
        db = get_db()
        try:
            row = db.execute("SELECT revoked_at FROM access_codes WHERE id = ?", (code_id,)).fetchone()
        except Exception:
            # access_codes table doesn't exist yet -- can only happen on the
            # very first request of a fresh deploy, before ensure_initialized()
            # has run init_db(). Let the request through rather than crash;
            # ensure_initialized() (registered below) fixes this immediately after.
            row = True
        db.close()
        if not row or (row is not True and row["revoked_at"] is not None):
            session.clear()
            return redirect(url_for("login"))


@app.route("/health")
def health():
    return jsonify({"status": "ok"})


@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        password = request.form.get("password", "")
        db = get_db()
        candidates = db.execute("SELECT * FROM access_codes WHERE revoked_at IS NULL").fetchall()
        db.close()
        matched = next((c for c in candidates if check_password_hash(c["password_hash"], password)), None)
        if matched:
            session["logged_in"] = True
            session["access_code_id"] = matched["id"]
            session["access_code_label"] = matched["label"]
            session["is_admin"] = bool(matched["is_admin"])
            return redirect(url_for("dashboard"))
        error = "Incorrect password — try again."
    return render_template("login.html", error=error)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


# ── Access codes (admin-only) ────────────────────────────────────────────────
def _require_admin():
    return session.get("is_admin") is True


@app.route("/settings/access-codes")
def access_codes_view():
    if not _require_admin():
        return render_template("error.html", message="This page is for the admin code only."), 403
    db = get_db()
    codes = db.execute("SELECT * FROM access_codes ORDER BY revoked_at IS NOT NULL, created_at").fetchall()
    active_admin_count = db.execute(
        "SELECT COUNT(*) AS n FROM access_codes WHERE is_admin = 1 AND revoked_at IS NULL"
    ).fetchone()["n"]
    db.close()
    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError:
        clients = []
    return render_template(
        "access_codes.html",
        clients=clients,
        active_client=clients[0] if clients else None,
        codes=codes,
        active_admin_count=active_admin_count,
        current_code_id=session.get("access_code_id"),
    )


@app.route("/settings/access-codes/new", methods=["POST"])
def access_codes_new():
    if not _require_admin():
        return render_template("error.html", message="This page is for the admin code only."), 403
    label = (request.form.get("label") or "").strip()
    password = request.form.get("password") or ""
    is_admin = 1 if request.form.get("is_admin") == "on" else 0
    if not label or not password:
        return render_template("error.html", message="Both a label and a password are required."), 400
    db = get_db()
    db.execute(
        "INSERT INTO access_codes (label, password_hash, is_admin) VALUES (?, ?, ?)",
        (label, generate_password_hash(password), is_admin)
    )
    db.commit()
    db.close()
    return redirect(url_for("access_codes_view"))


@app.route("/settings/access-codes/<int:code_id>/revoke", methods=["POST"])
def access_codes_revoke(code_id):
    if not _require_admin():
        return render_template("error.html", message="This page is for the admin code only."), 403
    db = get_db()
    row = db.execute("SELECT * FROM access_codes WHERE id = ?", (code_id,)).fetchone()
    if not row:
        db.close()
        return redirect(url_for("access_codes_view"))
    if row["is_admin"]:
        active_admin_count = db.execute(
            "SELECT COUNT(*) AS n FROM access_codes WHERE is_admin = 1 AND revoked_at IS NULL"
        ).fetchone()["n"]
        if active_admin_count <= 1:
            db.close()
            return render_template("error.html", message="Can't revoke the last active admin code -- you'd lock yourself out. Add another admin code first."), 400
    db.execute("UPDATE access_codes SET revoked_at = CURRENT_TIMESTAMP WHERE id = ?", (code_id,))
    db.commit()
    db.close()
    return redirect(url_for("access_codes_view"))


@app.route("/settings/access-codes/<int:code_id>/unrevoke", methods=["POST"])
def access_codes_unrevoke(code_id):
    """Undo an accidental revoke -- doesn't restore old sessions (those are
    gone), but the same password works again immediately."""
    if not _require_admin():
        return render_template("error.html", message="This page is for the admin code only."), 403
    db = get_db()
    db.execute("UPDATE access_codes SET revoked_at = NULL WHERE id = ?", (code_id,))
    db.commit()
    db.close()
    return redirect(url_for("access_codes_view"))


# ── Clients ───────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return redirect(url_for("dashboard"))


@app.route("/dashboard")
def dashboard():
    # Corrected 7/17: no local clients table -- Studio calls Hemingway's
    # real /api/clients directly (hemingway_client.py) as the single
    # source of truth, rather than keeping a copy that can drift.
    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502

    active_client_id = request.args.get("client_id", type=int)
    show_archived = request.args.get("show_archived") == "1"
    active_client = None
    projects = []
    degas_error = None
    if clients:
        active_client = next(
            (c for c in clients if c["id"] == active_client_id),
            clients[0]
        )
        db = get_db()
        if show_archived:
            rows = db.execute(
                "SELECT * FROM projects WHERE client_id = ? ORDER BY created_at DESC",
                (active_client["id"],)
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT * FROM projects WHERE client_id = ? AND archived_at IS NULL ORDER BY created_at DESC",
                (active_client["id"],)
            ).fetchall()

        # Task #7: Intake-through-Clipped phase comes from Degas's own
        # clips.status, read live rather than duplicated -- see
        # STUDIO_SYSTEM_DESIGN.md Section 4 "major simplification."
        # Drafting/Post Review are Studio-owned state (degas_project_id may
        # also be null for projects created before this was wired up).
        for p in rows:
            phase = p["phase"] or "intake"
            clips = []
            if p["degas_project_id"]:
                try:
                    degas_proj = degas_client.get_project(p["degas_project_id"])
                    clips = degas_proj.get("clips", [])
                    phase = degas_client.effective_phase(p["phase"], clips)
                except degas_client.DegasError as e:
                    degas_error = str(e)
            if phase != p["phase"]:
                db.execute("UPDATE projects SET phase = ? WHERE id = ?", (phase, p["id"]))
                db.commit()
            review_needed = db.execute(
                "SELECT COUNT(*) AS n FROM clip_review_flags WHERE project_id = ? "
                "AND (review_transcript = 1 OR review_video = 1)",
                (p["id"],)
            ).fetchone()["n"]
            projects.append({
                "id": p["id"],
                "name": p["name"],
                "degas_project_id": p["degas_project_id"],
                "phase": phase,
                "clip_count": len(clips),
                "clips_exported": sum(1 for c in clips if c["status"] == "exported"),
                "created_at": p["created_at"],
                "archived_at": p["archived_at"],
                "review_needed": review_needed,
            })
        db.close()

    cleanup_candidates, _cleanup_errors = _get_cleanup_candidates()

    return render_template(
        "dashboard.html",
        clients=clients,
        active_client=active_client,
        projects=projects,
        degas_error=degas_error,
        phase_labels=PHASE_LABELS,
        show_archived=show_archived,
        cleanup_count=len(cleanup_candidates),
    )


# ── Storage cleanup (task #22) ───────────────────────────────────────────────
# Nothing here ever deletes automatically in the background -- this only
# flags clips whose project is 45+ days old for Ben to review and delete
# himself from the Storage Cleanup page (or dismiss via the dashboard
# banner). Age is based on the Studio project's own created_at rather than
# asking Degas for per-clip timestamps, since Degas's existing JSON project
# endpoint doesn't return clip-level created_at and adding that would have
# meant a riskier hand-edit to an existing live route instead of a pure
# append -- see degas_client.delete_clip_media()'s docstring.
CLEANUP_THRESHOLD_DAYS = 45

# Chunk size for Studio's own server-to-server Nextcloud -> Degas import
# (task #28) -- matches the browser uploader's CHUNK_SIZE (project_detail.html)
# rather than reinventing a number, since that size is the one already proven
# to clear Degas's own upload path without issues.
NEXTCLOUD_IMPORT_CHUNK_SIZE = 8 * 1024 * 1024


def _get_cleanup_candidates():
    """Cross-client scan: every non-deleted clip belonging to a Studio
    project that's 45+ days old. Computed fresh on every call -- no caching
    yet, since project/clip counts are small today; worth revisiting if this
    gets slow once all 4 clients are onboarded (task #15)."""
    db = get_db()
    rows = db.execute("SELECT * FROM projects WHERE degas_project_id IS NOT NULL").fetchall()
    db.close()

    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError:
        clients = []
    client_names = {c["id"]: c["name"] for c in clients}

    now = datetime.now(timezone.utc)
    cutoff = now - timedelta(days=CLEANUP_THRESHOLD_DAYS)
    candidates = []
    errors = []
    for proj in rows:
        try:
            created_at = datetime.strptime(proj["created_at"], "%Y-%m-%d %H:%M:%S").replace(tzinfo=timezone.utc)
        except (ValueError, TypeError):
            continue
        if created_at >= cutoff:
            continue
        try:
            degas_proj = degas_client.get_project(proj["degas_project_id"])
        except degas_client.DegasError as e:
            errors.append(f"{proj['name']}: {e}")
            continue
        age_days = (now - created_at).days
        for clip in degas_proj.get("clips", []):
            if clip["status"] == "deleted":
                continue
            candidates.append({
                "project_id": proj["id"],
                "degas_project_id": proj["degas_project_id"],
                "project_name": proj["name"],
                "client_name": client_names.get(proj["client_id"], "?"),
                "clip_id": clip["id"],
                "filename": clip.get("original_filename") or clip["filename"],
                "status": clip["status"],
                "age_days": age_days,
            })
    return candidates, errors


@app.route("/settings/storage-cleanup")
def storage_cleanup_view():
    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502
    active_client = clients[0] if clients else None

    candidates, errors = _get_cleanup_candidates()
    return render_template(
        "storage_cleanup.html",
        clients=clients,
        active_client=active_client,
        candidates=candidates,
        errors=errors,
        threshold_days=CLEANUP_THRESHOLD_DAYS,
    )


@app.route("/settings/storage-cleanup/delete", methods=["POST"])
def storage_cleanup_delete():
    """Deletes exactly the clips Ben checked and confirmed -- nothing else.
    Each is gone from Degas's disk permanently; there's no undo."""
    selected = request.form.getlist("delete_ids")
    errors = []
    for item in selected:
        try:
            degas_project_id_str, clip_id_str = item.split(":")
            degas_client.delete_clip_media(int(degas_project_id_str), int(clip_id_str))
        except (ValueError, degas_client.DegasError) as e:
            errors.append(str(e))
    if errors:
        return render_template("error.html", message="; ".join(errors[:3])), 502
    return redirect(url_for("storage_cleanup_view"))


PHASE_LABELS = {
    "intake":         "Intake",
    "transcribing":   "Transcribing",
    "caption_review": "Caption Review",
    "clipped":        "Clipped",
    "drafting":       "Drafting",
    "post_review":    "Post Review",
}


@app.route("/clients/new", methods=["POST"])
def new_client():
    """Add a client without leaving Studio. Creates it in Hemingway (the
    single source of truth -- see hemingway_client.get_clients()), so it
    shows up there immediately with no separate sync step needed."""
    name = request.form.get("name", "").strip()
    if not name:
        return redirect(url_for("dashboard"))
    try:
        client = hemingway_client.create_client(name)
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502
    return redirect(url_for("dashboard", client_id=client["id"]))


@app.route("/projects/new", methods=["POST"])
def new_project():
    name = request.form.get("name", "").strip()
    client_id = request.form.get("client_id", type=int)
    if not name or not client_id:
        return redirect(url_for("dashboard", client_id=client_id))

    try:
        clients = hemingway_client.get_clients()
        client = next((c for c in clients if c["id"] == client_id), None)
        assigned_to = client["name"] if client else ""
        degas_project_id = degas_client.create_project(name, assigned_to, client_id)
    except (hemingway_client.HemingwayError, degas_client.DegasError) as e:
        return render_template("error.html", message=str(e)), 502

    db = get_db()
    db.execute(
        "INSERT INTO projects (client_id, name, degas_project_id, phase) VALUES (?, ?, ?, 'intake')",
        (client_id, name, degas_project_id)
    )
    db.commit()
    db.close()
    return redirect(url_for("dashboard", client_id=client_id))


@app.route("/projects/<int:project_id>/advance-phase", methods=["POST"])
def advance_phase(project_id):
    """Manual phase advance for Drafting/Post Review -- these are genuinely
    new state Studio owns (Section 4), not derived from Degas, so they need
    an explicit action rather than being inferred."""
    target = request.form.get("target")
    client_id = request.form.get("client_id", type=int)
    allowed_transitions = {"clipped": "drafting", "drafting": "post_review"}

    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    if proj and allowed_transitions.get(proj["phase"]) == target:
        db.execute("UPDATE projects SET phase = ? WHERE id = ?", (target, project_id))
        db.commit()
    db.close()
    return redirect(url_for("dashboard", client_id=client_id))


@app.route("/projects/<int:project_id>/archive", methods=["POST"])
def project_archive(project_id):
    client_id = request.form.get("client_id", type=int)
    db = get_db()
    db.execute("UPDATE projects SET archived_at = CURRENT_TIMESTAMP WHERE id = ?", (project_id,))
    db.commit()
    db.close()
    return redirect(url_for("dashboard", client_id=client_id))


@app.route("/projects/<int:project_id>/unarchive", methods=["POST"])
def project_unarchive(project_id):
    client_id = request.form.get("client_id", type=int)
    db = get_db()
    db.execute("UPDATE projects SET archived_at = NULL WHERE id = ?", (project_id,))
    db.commit()
    db.close()
    return redirect(url_for("dashboard", client_id=client_id))


@app.route("/projects/<int:project_id>/delete", methods=["POST"])
def project_delete(project_id):
    """Permanently removes this project from Studio only -- does not touch
    the linked Degas project or its uploaded video/exported files, which are
    a separate system Studio doesn't own and shouldn't destroy from here.
    Posts already generated/scheduled keep existing (posts.project_id ON
    DELETE SET NULL) -- deleting the project just unlinks them, it doesn't
    delete post history."""
    client_id = request.form.get("client_id", type=int)
    db = get_db()
    db.execute("DELETE FROM projects WHERE id = ?", (project_id,))
    db.commit()
    db.close()
    return redirect(url_for("dashboard", client_id=client_id))


# ── Project workspace (task #21): the real Upload -> Transcribe -> Review ->
# Export -> Write Posts pipeline, built as genuine Studio screens against
# Degas's and Hemingway's real APIs, not stubs or redirects to those other
# apps -- confirmed against live server source 7/24, no changes needed on
# either Degas's or Hemingway's side.
def _build_project_transcript(degas_project_id, clips):
    """Builds a transcript string compatible with Hemingway's
    split_transcript() parser: one 'VIDEO: NN - Title' line per clip
    followed by that clip's current (possibly human-edited) segment text,
    one segment per line. Only includes clips whose transcript actually
    exists yet (transcribed/exported) -- clips still uploading/transcribing/
    erroring have nothing usable.

    Returns (transcript_str, ordered_clip_ids) -- ordered_clip_ids[i] is the
    Degas clip id behind the i-th section actually included in the
    transcript. Hemingway's /api/generate streams back posts with a
    positional 'index' matching split_transcript()'s section order, so
    callers can zip that index against this list to know exactly which clip
    a generated post's copy came from (task #27, Ben's ask 7/24: "I want the
    copy to connect to its correlating video")."""
    sections = []
    ordered_clip_ids = []
    eligible = [c for c in clips if c["status"] in ("transcribed", "exported")]
    for idx, clip in enumerate(eligible, start=1):
        title = os.path.splitext(clip.get("original_filename") or clip["filename"])[0]
        try:
            seg_data = degas_client.get_clip_segments(degas_project_id, clip["id"])
        except degas_client.DegasError:
            continue
        segments = seg_data.get("current") or seg_data.get("original") or []
        lines = [f"VIDEO: {idx:02d} - {title}"]
        for seg in segments:
            text = (seg.get("text") or "").strip()
            if text:
                lines.append(text)
        if len(lines) > 1:
            sections.append("\n".join(lines))
            ordered_clip_ids.append(clip["id"])
    return "\n\n".join(sections), ordered_clip_ids


@app.route("/projects/<int:project_id>")
def project_detail(project_id):
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    if not proj:
        db.close()
        return redirect(url_for("dashboard"))

    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        db.close()
        return render_template("error.html", message=str(e)), 502
    active_client = next((c for c in clients if c["id"] == proj["client_id"]), None)

    degas_clips = []
    degas_error = None
    if proj["degas_project_id"]:
        try:
            degas_proj = degas_client.get_project(proj["degas_project_id"])
            degas_clips = degas_proj.get("clips", [])
        except degas_client.DegasError as e:
            degas_error = str(e)

    project_posts = db.execute(
        "SELECT * FROM posts WHERE project_id = ? ORDER BY created_at DESC", (project_id,)
    ).fetchall()
    linked = db.execute(
        "SELECT * FROM client_postiz_groups WHERE client_id = ?", (proj["client_id"],)
    ).fetchone()
    db.close()

    # Media-capable set (not the text-only default) so YouTube/Instagram are
    # real options here, same as the calendar's Add Post modal (task #18) --
    # except here the video comes from the post's own linked clip rather than
    # a manual upload (task #27, Ben's ask 7/24).
    channels, channels_error = _get_schedulable_channels(linked, postiz_client.MEDIA_CAPABLE_IDENTIFIERS)
    can_write_posts = any(c["status"] in ("transcribed", "exported") for c in degas_clips)

    clip_by_id = {c["id"]: c for c in degas_clips}
    posts_view = []
    for p in project_posts:
        pd = dict(p)
        clip = clip_by_id.get(pd.get("clip_id"))
        pd["clip_exported"] = bool(clip and clip["status"] == "exported")
        title_source = clip.get("original_filename") or clip.get("filename") if clip else None
        pd["suggested_youtube_title"] = os.path.splitext(title_source)[0] if title_source else proj["name"]
        posts_view.append(pd)

    nextcloud_folder, nextcloud_files, nextcloud_error = _get_nextcloud_incoming_files(proj["client_id"])

    return render_template(
        "project_detail.html",
        clients=clients,
        active_client=active_client,
        project=proj,
        phase_labels=PHASE_LABELS,
        degas_clips=degas_clips,
        degas_error=degas_error,
        export_styles=degas_client.EXPORT_STYLES,
        project_posts=posts_view,
        can_write_posts=can_write_posts,
        linked=linked,
        channels=channels,
        channels_error=channels_error,
        media_required_identifiers=list(postiz_client.MEDIA_REQUIRED_IDENTIFIERS),
        nextcloud_folder=nextcloud_folder,
        nextcloud_files=nextcloud_files,
        nextcloud_error=nextcloud_error,
    )


def _archive_direct_upload_to_cloud(proj, filename):
    """Best-effort copy of a freshly-completed direct upload (Studio's own
    Upload button, not the Cloud KMG pull-in path) out to /imported, so
    files land there regardless of which upload path someone used (task
    #31, Ben's ask 8/11: 'Yes, I want that'). Matches the clip by filename
    since Degas's chunk-upload response doesn't return the new clip's id --
    picks the last match if there happen to be several same-named clips,
    since that's the one most likely to be the one that just finished.

    Non-fatal by design: the upload into Degas already succeeded by the
    time this runs, which is what actually matters for the pipeline to
    keep working -- a failed Nextcloud copy shouldn't turn a successful
    upload into an error for the browser's uploader."""
    linked = _get_nextcloud_folder(proj["client_id"])
    if not linked:
        return
    try:
        degas_proj = degas_client.get_project(proj["degas_project_id"])
        matches = [
            c for c in degas_proj.get("clips", [])
            if (c.get("original_filename") or c.get("filename")) == filename
        ]
        if not matches:
            return
        clip = matches[-1]
        video_resp = degas_client.stream_clip_video(proj["degas_project_id"], clip["id"])
        nextcloud_client.ensure_folder(f"{linked['folder_name']}/imported")
        nextcloud_client.upload_file(
            f"{linked['folder_name']}/imported/{filename}", video_resp.content, "video/mp4"
        )
    except (degas_client.DegasError, nextcloud_client.NextcloudError):
        pass


@app.route("/projects/<int:project_id>/upload-chunk", methods=["POST"])
def project_upload_chunk(project_id):
    """Receives one chunk from the browser's JS uploader and forwards it to
    Degas's real chunked-upload route using Studio's own authenticated
    session -- the browser never talks to degas.kmgtools.us directly."""
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj or not proj["degas_project_id"]:
        return jsonify({"error": "This project isn't linked to Degas."}), 400

    file_uid = request.form.get("file_uid", "")
    chunk_index = request.form.get("chunk_index", "")
    total_chunks = request.form.get("total_chunks", "")
    filename = request.form.get("filename", "")
    file_obj = request.files.get("data")
    if not file_uid or not file_obj:
        return jsonify({"error": "missing fields"}), 400

    try:
        result = degas_client.upload_chunk(
            proj["degas_project_id"], file_uid, chunk_index, total_chunks, filename,
            file_obj.stream.read(), file_obj.content_type or "application/octet-stream",
        )
    except degas_client.DegasError as e:
        return jsonify({"error": str(e)}), 502

    if result.get("status") == "complete":
        _archive_direct_upload_to_cloud(proj, filename)

    return jsonify(result)


@app.route("/projects/<int:project_id>/nextcloud-import", methods=["POST"])
def project_nextcloud_import(project_id):
    """Pulls one file from the client's Cloud KMG /incoming folder straight
    into this project's Degas upload pipeline -- task #28, Ben's ask 8/3:
    'remote team members can work on projects without moving files to
    personal computers'. Downloads the file from Nextcloud, then re-chunks
    and forwards it through Degas's existing chunked-upload route exactly
    the way the browser's own JS uploader does (project_detail.html), just
    sourced from Nextcloud bytes instead of a local <input type=file>."""
    filename = request.form.get("filename", "").strip()
    if not filename:
        return render_template("error.html", message="No file selected to import."), 400

    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj or not proj["degas_project_id"]:
        return render_template("error.html", message="This project isn't linked to Degas."), 400

    linked = _get_nextcloud_folder(proj["client_id"])
    if not linked:
        return render_template("error.html", message="This client isn't linked to a Cloud KMG folder yet -- set that up first."), 400

    try:
        nc_resp = nextcloud_client.download_file(f"{linked['folder_name']}/incoming/{filename}")
        file_bytes = nc_resp.content
    except nextcloud_client.NextcloudError as e:
        return render_template("error.html", message=str(e)), 502

    if not file_bytes:
        return render_template("error.html", message=f"'{filename}' downloaded empty from Cloud KMG -- nothing to import."), 400

    file_uid = uuid.uuid4().hex
    total_chunks = max(1, (len(file_bytes) + NEXTCLOUD_IMPORT_CHUNK_SIZE - 1) // NEXTCLOUD_IMPORT_CHUNK_SIZE)
    try:
        for i in range(total_chunks):
            start = i * NEXTCLOUD_IMPORT_CHUNK_SIZE
            chunk = file_bytes[start:start + NEXTCLOUD_IMPORT_CHUNK_SIZE]
            degas_client.upload_chunk(
                proj["degas_project_id"], file_uid, i, total_chunks, filename,
                chunk, "application/octet-stream",
            )
    except degas_client.DegasError as e:
        return render_template("error.html", message=f"Import failed partway through: {e}"), 502

    # Move the source file out of /incoming so it doesn't keep showing up as
    # 'unclaimed' footage once it's actually been pulled into a project
    # (task #30, Ben's ask 8/11). Best-effort: the import into Degas already
    # succeeded above, which is the part that actually matters -- a failed
    # tidy-up here shouldn't turn a successful import into an error page.
    try:
        nextcloud_client.ensure_folder(f"{linked['folder_name']}/imported")
        nextcloud_client.move_file(
            f"{linked['folder_name']}/incoming/{filename}",
            f"{linked['folder_name']}/imported/{filename}",
        )
    except nextcloud_client.NextcloudError:
        pass

    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/transcribe", methods=["POST"])
def clip_transcribe(project_id, clip_id):
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if proj and proj["degas_project_id"]:
        try:
            degas_client.trigger_transcribe(proj["degas_project_id"], clip_id)
        except degas_client.DegasError as e:
            return render_template("error.html", message=str(e)), 502
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/transcribe-all", methods=["POST"])
def project_transcribe_all(project_id):
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if proj and proj["degas_project_id"]:
        try:
            degas_client.trigger_transcribe_all(proj["degas_project_id"])
        except degas_client.DegasError as e:
            return render_template("error.html", message=str(e)), 502
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/status.json")
def clip_status_json(project_id, clip_id):
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj or not proj["degas_project_id"]:
        return jsonify({"error": "not linked"}), 400
    try:
        return jsonify(degas_client.get_clip_status(proj["degas_project_id"], clip_id))
    except degas_client.DegasError as e:
        return jsonify({"error": str(e)}), 502


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/review")
def clip_review(project_id, clip_id):
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj:
        return redirect(url_for("dashboard"))

    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502
    active_client = next((c for c in clients if c["id"] == proj["client_id"]), None)

    try:
        seg_data = degas_client.get_clip_segments(proj["degas_project_id"], clip_id)
    except degas_client.DegasError as e:
        return render_template("error.html", message=str(e)), 502

    flags = _get_review_flags(project_id, [clip_id])[clip_id]

    return render_template(
        "clip_review.html",
        clients=clients,
        active_client=active_client,
        project=proj,
        clip_id=clip_id,
        segments=seg_data.get("current") or [],
        review_flags=flags,
    )


def _get_review_flags(project_id, clip_ids):
    """Returns {clip_id: {'review_transcript': bool, 'review_video': bool}}
    for the given clips. Clips with no row default to both False (normal)."""
    if not clip_ids:
        return {}
    db = get_db()
    placeholders = ",".join("?" * len(clip_ids))
    rows = db.execute(
        f"SELECT clip_id, review_transcript, review_video FROM clip_review_flags "
        f"WHERE project_id = ? AND clip_id IN ({placeholders})",
        (project_id, *clip_ids)
    ).fetchall()
    db.close()
    flags = {r["clip_id"]: {"review_transcript": bool(r["review_transcript"]), "review_video": bool(r["review_video"])} for r in rows}
    for cid in clip_ids:
        flags.setdefault(cid, {"review_transcript": False, "review_video": False})
    return flags


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/review-flags", methods=["POST"])
def clip_review_flags_update(project_id, clip_id):
    """Toggle the Review Transcript (orange) / Review Video (purple) flags
    for one clip. Checkboxes auto-submit their own tiny form on change
    (see clip_review.html/review_all.html), so this always receives the
    FULL desired state of both checkboxes, not a single toggle."""
    review_transcript = 1 if request.form.get("review_transcript") == "on" else 0
    review_video = 1 if request.form.get("review_video") == "on" else 0
    db = get_db()
    db.execute(
        """INSERT INTO clip_review_flags (project_id, clip_id, review_transcript, review_video, updated_at)
           VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
           ON CONFLICT(project_id, clip_id) DO UPDATE SET
             review_transcript = excluded.review_transcript,
             review_video = excluded.review_video,
             updated_at = CURRENT_TIMESTAMP""",
        (project_id, clip_id, review_transcript, review_video)
    )
    db.commit()
    db.close()
    return _review_flags_redirect(project_id, clip_id)


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/review-flags/approve", methods=["POST"])
def clip_review_flags_approve(project_id, clip_id):
    """'Approved' means the flagged issue was addressed -- clears both flags
    back to normal rather than being a third independent flag. The frontend
    already disables this checkbox unless one of the other two is set, but
    this route doesn't re-check that -- clearing an already-clear row is a
    harmless no-op."""
    db = get_db()
    db.execute(
        "UPDATE clip_review_flags SET review_transcript = 0, review_video = 0, updated_at = CURRENT_TIMESTAMP "
        "WHERE project_id = ? AND clip_id = ?",
        (project_id, clip_id)
    )
    db.commit()
    db.close()
    return _review_flags_redirect(project_id, clip_id)


def _review_flags_redirect(project_id, clip_id):
    """Send the user back to whichever screen they came from (single-clip
    Review vs. Review All), landing on this clip so the page doesn't jump
    around after a checkbox click."""
    came_from = request.form.get("came_from")
    if came_from == "review_all":
        return redirect(url_for("project_review_all", project_id=project_id) + f"#clip-{clip_id}")
    return redirect(url_for("clip_review", project_id=project_id, clip_id=clip_id))


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/review/save", methods=["POST"])
def clip_review_save(project_id, clip_id):
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj:
        return redirect(url_for("dashboard"))

    try:
        seg_data = degas_client.get_clip_segments(proj["degas_project_id"], clip_id)
    except degas_client.DegasError as e:
        return render_template("error.html", message=str(e)), 502

    current = seg_data.get("current") or []
    texts = request.form.getlist("segment_text")
    merged = []
    for i, seg in enumerate(current):
        text = texts[i] if i < len(texts) else seg.get("text", "")
        merged.append({"start": seg.get("start", 0), "end": seg.get("end", 0), "text": text})

    try:
        degas_client.save_clip_segments(proj["degas_project_id"], clip_id, merged)
    except degas_client.DegasError as e:
        return render_template("error.html", message=str(e)), 502

    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/review-all")
def project_review_all(project_id):
    """Every transcribed/exported clip's transcript in one view, matching
    Degas's own editor-all page (Ben's ask, 7/24) -- built as a genuine
    Studio screen against Degas's real per-clip segments route + its bulk
    /save-all route, not a redirect to Degas's own UI."""
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj:
        return redirect(url_for("dashboard"))

    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502
    active_client = next((c for c in clients if c["id"] == proj["client_id"]), None)

    clips_data = []
    degas_error = None
    if proj["degas_project_id"]:
        try:
            degas_proj = degas_client.get_project(proj["degas_project_id"])
            eligible = [c for c in degas_proj.get("clips", []) if c["status"] in ("transcribed", "exported")]
            for clip in eligible:
                try:
                    seg_data = degas_client.get_clip_segments(proj["degas_project_id"], clip["id"])
                except degas_client.DegasError as e:
                    degas_error = str(e)
                    continue
                clips_data.append({
                    "clip_id": clip["id"],
                    "filename": clip.get("original_filename") or clip["filename"],
                    "segments": seg_data.get("current") or [],
                })
        except degas_client.DegasError as e:
            degas_error = str(e)

    review_flags = _get_review_flags(project_id, [c["clip_id"] for c in clips_data])

    return render_template(
        "review_all.html",
        clients=clients,
        active_client=active_client,
        project=proj,
        clips_data=clips_data,
        degas_error=degas_error,
        review_flags=review_flags,
    )


@app.route("/projects/<int:project_id>/review-all/save", methods=["POST"])
def project_review_all_save(project_id):
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj or not proj["degas_project_id"]:
        return redirect(url_for("dashboard"))

    clip_ids = request.form.getlist("clip_id")
    payload = []
    for clip_id_str in clip_ids:
        clip_id = int(clip_id_str)
        try:
            seg_data = degas_client.get_clip_segments(proj["degas_project_id"], clip_id)
        except degas_client.DegasError as e:
            return render_template("error.html", message=str(e)), 502
        current = seg_data.get("current") or []
        texts = request.form.getlist(f"segment_text_{clip_id}")
        merged = []
        for i, seg in enumerate(current):
            text = texts[i] if i < len(texts) else seg.get("text", "")
            merged.append({"start": seg.get("start", 0), "end": seg.get("end", 0), "text": text})
        payload.append({"clip_id": clip_id, "segments": merged})

    if payload:
        try:
            degas_client.save_all_clip_segments(proj["degas_project_id"], payload)
        except degas_client.DegasError as e:
            return render_template("error.html", message=str(e)), 502

    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/export", methods=["POST"])
def clip_export(project_id, clip_id):
    style = request.form.get("style", "1")
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if proj and proj["degas_project_id"]:
        try:
            degas_client.trigger_export(proj["degas_project_id"], clip_id, style)
        except degas_client.DegasError as e:
            return render_template("error.html", message=str(e)), 502
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/export-all", methods=["POST"])
def project_export_all(project_id):
    style = request.form.get("style", "1")
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if proj and proj["degas_project_id"]:
        try:
            degas_client.trigger_export_all(proj["degas_project_id"], style)
        except degas_client.DegasError as e:
            return render_template("error.html", message=str(e)), 502
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/download")
def clip_download(project_id, clip_id):
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj or not proj["degas_project_id"]:
        return render_template("error.html", message="This project isn't linked to Degas."), 400

    try:
        degas_resp = degas_client.download_clip(proj["degas_project_id"], clip_id)
    except degas_client.DegasError as e:
        return render_template("error.html", message=str(e)), 502

    content_disposition = degas_resp.headers.get("Content-Disposition")
    if not content_disposition:
        return render_template("error.html", message="This clip isn't exported yet -- export it first, then download."), 400

    return Response(
        degas_resp.iter_content(chunk_size=8192),
        content_type=degas_resp.headers.get("Content-Type", "video/mp4"),
        headers={"Content-Disposition": content_disposition},
    )


@app.route("/projects/<int:project_id>/clips/download-all")
def project_download_all(project_id):
    """Zips every exported clip in the project and sends it as one download.
    Built to a temp file on disk (not buffered in memory) since projects can
    have several large video files -- this server doesn't have RAM to spare
    for holding a multi-hundred-MB zip in memory (see transcription.py's
    accuracy-history comment for exactly how tight that got once already).
    Best-effort per clip: a single clip failing to download from Degas skips
    that clip rather than failing the whole zip."""
    import zipfile
    import tempfile

    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj or not proj["degas_project_id"]:
        return render_template("error.html", message="This project isn't linked to Degas."), 400

    try:
        degas_proj = degas_client.get_project(proj["degas_project_id"])
    except degas_client.DegasError as e:
        return render_template("error.html", message=str(e)), 502

    exported_clips = [c for c in degas_proj.get("clips", []) if c.get("status") == "exported"]
    if not exported_clips:
        return render_template("error.html", message="No exported clips to download yet."), 400

    fd, zip_path = tempfile.mkstemp(suffix=".zip")
    os.close(fd)

    used_names = set()
    any_written = False
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_STORED) as zf:
        for c in exported_clips:
            try:
                resp = degas_client.download_clip(proj["degas_project_id"], c["id"])
            except degas_client.DegasError:
                continue  # best-effort -- skip a clip that fails rather than failing the whole zip
            if not resp.headers.get("Content-Disposition"):
                continue

            base_name = c.get("original_filename") or c.get("filename") or f"clip-{c['id']}.mp4"
            name = base_name
            n = 1
            while name in used_names:
                stem, dot, ext = base_name.rpartition(".")
                name = f"{stem} ({n}){dot}{ext}" if dot else f"{base_name} ({n})"
                n += 1
            used_names.add(name)

            with zf.open(name, "w") as entry:
                for chunk in resp.iter_content(chunk_size=8192):
                    if chunk:
                        entry.write(chunk)
            any_written = True

    if not any_written:
        os.remove(zip_path)
        return render_template("error.html", message="Couldn't reach Degas to download any clips -- try again."), 502

    project_name = (proj["name"] or f"project-{project_id}").strip()
    download_name = f"{project_name} - clips.zip"

    @after_this_request
    def _cleanup(response):
        try:
            os.remove(zip_path)
        except OSError:
            pass
        return response

    return send_file(zip_path, as_attachment=True, download_name=download_name, mimetype="application/zip")


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/archive-to-cloud", methods=["POST"])
def clip_archive_to_cloud(project_id, clip_id):
    """Pushes a clip's finished EXPORTED (captioned) video out to the
    client's Cloud KMG /captioned folder -- task #28, the other half of
    Ben's ask (8/3): remote team members can grab finished clips without
    Studio being the only place they exist. Manual, one click per clip
    (same pattern as the existing Download button) rather than automatic
    right after export -- Studio has no reliable server-side signal for
    'export just finished', only the browser's own polling JS discovers
    that (see pollStatuses() in project_detail.html)."""
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj or not proj["degas_project_id"]:
        return render_template("error.html", message="This project isn't linked to Degas."), 400

    linked = _get_nextcloud_folder(proj["client_id"])
    if not linked:
        return render_template("error.html", message="This client isn't linked to a Cloud KMG folder yet -- set that up first."), 400

    try:
        degas_resp = degas_client.download_clip(proj["degas_project_id"], clip_id)
    except degas_client.DegasError as e:
        return render_template("error.html", message=str(e)), 502

    content_disposition = degas_resp.headers.get("Content-Disposition")
    if not content_disposition:
        return render_template("error.html", message="This clip isn't exported yet -- export it first, then archive."), 400

    filename = None
    if "filename=" in content_disposition:
        filename = content_disposition.split("filename=", 1)[1].strip('"; ')
    filename = filename or f"clip_{clip_id}.mp4"

    try:
        nextcloud_client.ensure_folder(f"{linked['folder_name']}/captioned")
        nextcloud_client.upload_file(
            f"{linked['folder_name']}/captioned/{filename}", degas_resp.content, "video/mp4"
        )
    except nextcloud_client.NextcloudError as e:
        return render_template("error.html", message=str(e)), 502

    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/clips/<int:clip_id>/video")
def clip_video(project_id, clip_id):
    """Proxies the clip's raw uploaded video (no captions) for inline
    preview next to the transcript during Caption Review -- so Ben can
    watch the clip and read/edit the text side by side to catch the class
    of error text-only review can't: Whisper hearing a name or number
    wrong but writing something clean and grammatical (Ben's ask, 7/24).

    Forwards the browser's Range header to Degas and mirrors back whatever
    status/headers Degas responds with, so the <video> player's seek bar
    actually works instead of only playing top-to-bottom."""
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    db.close()
    if not proj or not proj["degas_project_id"]:
        return render_template("error.html", message="This project isn't linked to Degas."), 400

    range_header = request.headers.get("Range")
    try:
        degas_resp = degas_client.stream_clip_video(proj["degas_project_id"], clip_id, range_header)
    except degas_client.DegasError as e:
        return render_template("error.html", message=str(e)), 502

    proxy_headers = {}
    for h in ("Content-Type", "Content-Length", "Content-Range", "Accept-Ranges"):
        if h in degas_resp.headers:
            proxy_headers[h] = degas_resp.headers[h]
    proxy_headers.setdefault("Content-Type", "video/mp4")
    proxy_headers.setdefault("Accept-Ranges", "bytes")

    return Response(
        degas_resp.iter_content(chunk_size=65536),
        status=degas_resp.status_code,
        headers=proxy_headers,
    )


@app.route("/projects/<int:project_id>/write-posts", methods=["POST"])
def project_write_posts(project_id):
    """Generates real posts from this project's actual reviewed transcript --
    the counterpart to Quick Posts (task #11), which deliberately has no
    Degas transcript involved. This is the 'Write Captions in Hemingway
    Module' step of the pipeline Ben asked for (7/24)."""
    style = request.form.get("style", "conversational")
    length = request.form.get("length", "short")
    context = request.form.get("context", "").strip()

    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    if not proj:
        db.close()
        return render_template("error.html", message="Project not found."), 404
    if not proj["degas_project_id"]:
        db.close()
        return render_template("error.html", message="This project isn't linked to Degas -- nothing to write from."), 400

    try:
        degas_proj = degas_client.get_project(proj["degas_project_id"])
    except degas_client.DegasError as e:
        db.close()
        return render_template("error.html", message=str(e)), 502

    transcript, ordered_clip_ids = _build_project_transcript(proj["degas_project_id"], degas_proj.get("clips", []))
    if not transcript.strip():
        db.close()
        return render_template("error.html", message="No reviewed transcript text yet -- transcribe and review at least one clip before writing posts."), 400

    try:
        result = hemingway_client.generate_from_transcript(
            proj["client_id"], transcript, style=style, length=length, context=context, name=proj["name"]
        )
    except hemingway_client.HemingwayError as e:
        db.close()
        return render_template("error.html", message=str(e)), 502

    db.execute(
        "UPDATE projects SET hemingway_batch_id = ?, phase = 'drafting' WHERE id = ?",
        (result["batch_id"], project_id)
    )
    for p in result["posts"]:
        if p.get("id") and p.get("body"):
            idx = p.get("index")
            clip_id = ordered_clip_ids[idx] if isinstance(idx, int) and 0 <= idx < len(ordered_clip_ids) else None
            db.execute(
                """INSERT INTO posts (client_id, project_id, source, caption, status, hemingway_post_id, clip_id, title)
                   VALUES (?, ?, 'project', ?, 'draft', ?, ?, ?)""",
                (proj["client_id"], project_id, p["body"], p["id"], clip_id, p.get("title"))
            )
    db.commit()
    db.close()
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/posts/export-doc")
def project_posts_export_doc(project_id):
    """One-click export of every post written for this project into a single
    .docx -- title (which video it came from) then the post copy, in the
    order they were generated. Ben's ask, 2026-08-26: a review/handoff doc
    he can send along without opening Studio."""
    db = get_db()
    proj = db.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    if not proj:
        db.close()
        return redirect(url_for("dashboard"))

    posts = db.execute(
        "SELECT * FROM posts WHERE project_id = ? ORDER BY id ASC", (project_id,)
    ).fetchall()
    db.close()

    if not posts:
        return render_template("error.html", message="No posts written yet for this project."), 400

    doc = Document()
    doc.add_heading(proj["name"] or f"Project {project_id}", level=1)

    for i, p in enumerate(posts):
        heading = p["title"] or f"Post {i + 1}"
        doc.add_heading(heading, level=2)
        body_para = doc.add_paragraph(p["caption"] or "")
        body_para.style.font.size = Pt(11)
        if i < len(posts) - 1:
            doc.add_paragraph("")  # spacer between posts

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    project_name = (proj["name"] or f"project-{project_id}").strip()
    download_name = f"{project_name} - posts.docx"

    return send_file(
        buf,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/clients")
def api_clients():
    try:
        return jsonify(hemingway_client.get_clients())
    except hemingway_client.HemingwayError as e:
        return jsonify({"error": str(e)}), 502


# ── Glossary (task #8) ───────────────────────────────────────────────────────
@app.route("/clients/<int:client_id>/glossary")
def glossary_view(client_id):
    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502
    active_client = next((c for c in clients if c["id"] == client_id), None)
    if not active_client:
        return redirect(url_for("dashboard"))

    db = get_db()
    pending = db.execute(
        "SELECT * FROM glossary_terms WHERE client_id = ? AND status = 'pending' ORDER BY occurrence_count DESC",
        (client_id,)
    ).fetchall()
    confirmed_client = db.execute(
        "SELECT * FROM glossary_terms WHERE client_id = ? AND status = 'confirmed' ORDER BY term",
        (client_id,)
    ).fetchall()
    confirmed_global = db.execute(
        "SELECT * FROM glossary_terms WHERE client_id IS NULL AND status = 'confirmed' ORDER BY term"
    ).fetchall()
    db.close()

    return render_template(
        "glossary.html",
        clients=clients,
        active_client=active_client,
        pending=pending,
        confirmed_client=confirmed_client,
        confirmed_global=confirmed_global,
    )


@app.route("/clients/<int:client_id>/glossary/scan", methods=["POST"])
def glossary_scan(client_id):
    """Scans every Degas-linked project for this client, diffs each clip's
    original vs. current transcript, and records any new candidates as
    pending. Pull-on-click, not a background job -- same reasoning as the
    Postiz analytics design (task #16): cheap enough to run when asked,
    no reason to poll constantly."""
    db = get_db()
    projects = db.execute(
        "SELECT * FROM projects WHERE client_id = ? AND degas_project_id IS NOT NULL",
        (client_id,)
    ).fetchall()

    total_new = 0
    errors = []
    for proj in projects:
        try:
            degas_proj = degas_client.get_project(proj["degas_project_id"])
        except degas_client.DegasError as e:
            errors.append(str(e))
            continue
        for clip in degas_proj.get("clips", []):
            try:
                seg_data = degas_client.get_clip_segments(proj["degas_project_id"], clip["id"])
            except degas_client.DegasError as e:
                errors.append(str(e))
                continue
            candidates = glossary.detect_candidates(seg_data.get("original", []), seg_data.get("current", []))
            total_new += glossary.record_candidates(db, client_id, candidates)
    db.close()

    if errors:
        return render_template("error.html", message="; ".join(errors[:3])), 502
    return redirect(url_for("glossary_view", client_id=client_id, found=total_new))


@app.route("/glossary/<int:term_id>/confirm", methods=["POST"])
def glossary_confirm(term_id):
    client_id = request.form.get("client_id", type=int)
    category = request.form.get("category", "").strip()
    db = get_db()
    if category:
        db.execute(
            "UPDATE glossary_terms SET status = 'confirmed', category = ? WHERE id = ?",
            (category, term_id)
        )
    else:
        db.execute("UPDATE glossary_terms SET status = 'confirmed' WHERE id = ?", (term_id,))
    db.commit()
    db.close()
    return redirect(url_for("glossary_view", client_id=client_id))


@app.route("/glossary/<int:term_id>/reject", methods=["POST"])
def glossary_reject(term_id):
    client_id = request.form.get("client_id", type=int)
    db = get_db()
    db.execute("DELETE FROM glossary_terms WHERE id = ? AND status = 'pending'", (term_id,))
    db.commit()
    db.close()
    return redirect(url_for("glossary_view", client_id=client_id))


@app.route("/glossary/<int:term_id>/promote", methods=["POST"])
def glossary_promote(term_id):
    """Hoists a client-specific confirmed term to global -- Section 4:
    'a "promote to global" action lets you hoist something from a client
    glossary once you notice it's not actually client-specific.'"""
    client_id = request.form.get("client_id", type=int)
    db = get_db()
    db.execute(
        "UPDATE glossary_terms SET client_id = NULL WHERE id = ? AND status = 'confirmed'",
        (term_id,)
    )
    db.commit()
    db.close()
    return redirect(url_for("glossary_view", client_id=client_id))


# ── Quick posts (task #11) ───────────────────────────────────────────────────
QUICK_POST_STYLES = ["thought-leader", "conversational", "storyteller", "punchy"]
QUICK_POST_LENGTHS = ["super-short", "short", "medium", "long"]
# draft -> scheduled now happens for real via quick_posts_schedule_to_postiz
# (task #13) -- Postiz itself is the source of truth for that transition.
# scheduled -> published stays a manual confirm for now since Studio doesn't
# sync Postiz's publish webhook/state back yet (the calendar shows Postiz's
# real state directly instead -- see calendar_view).
QUICK_POST_TRANSITIONS = {"scheduled": "published"}


def _post_redirect(post, client_id):
    """Quick Posts (task #11) and project-sourced posts (task #21) share the
    same posts table and the same edit/regenerate/schedule actions -- this
    sends you back to wherever that post actually lives (its project page,
    or Quick Posts) instead of hardcoding Quick Posts for both."""
    if post and post["project_id"]:
        return redirect(url_for("project_detail", project_id=post["project_id"]))
    return redirect(url_for("quick_posts_view", client_id=client_id))


def _get_schedulable_channels(linked, allowed_identifiers=None):
    """Shared by quick_posts_view and calendar_view -- returns (channels,
    channels_error) for whichever client's Postiz group is passed in,
    filtered to a set of identifiers Studio actually knows how to schedule
    to. Defaults to the text-only set (Quick Posts has no media upload UI);
    calendar_view passes MEDIA_CAPABLE_IDENTIFIERS since it has one. Factored
    out so both screens can't drift on what "available channels" means."""
    if not linked:
        return [], None
    allowed = allowed_identifiers or postiz_client.SUPPORTED_SCHEDULE_IDENTIFIERS
    try:
        all_integrations = postiz_client.list_integrations(linked["postiz_group_id"])
        channels = [
            i for i in all_integrations
            if i["identifier"] in allowed and not i.get("disabled")
        ]
        return channels, None
    except postiz_client.PostizError as e:
        return [], str(e)


def _get_nextcloud_folder(client_id):
    """The client's linked Cloud KMG folder row, or None if not set up yet."""
    db = get_db()
    linked = db.execute("SELECT * FROM client_nextcloud_folders WHERE client_id = ?", (client_id,)).fetchone()
    db.close()
    return linked


def _get_nextcloud_incoming_files(client_id):
    """Returns (folder_name, files, error) for a client's Cloud KMG
    /incoming subfolder (task #28). files is [] rather than an error both
    when nothing's linked yet and when the folder's simply empty -- callers
    tell those apart via folder_name being None."""
    linked = _get_nextcloud_folder(client_id)
    if not linked:
        return None, [], None
    try:
        files = nextcloud_client.list_folder(f"{linked['folder_name']}/incoming")
        return linked["folder_name"], files, None
    except nextcloud_client.NextcloudError as e:
        return linked["folder_name"], [], str(e)


@app.route("/clients/<int:client_id>/quick-posts")
def quick_posts_view(client_id):
    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502
    active_client = next((c for c in clients if c["id"] == client_id), None)
    if not active_client:
        return redirect(url_for("dashboard"))

    db = get_db()
    quick_posts = db.execute(
        "SELECT * FROM posts WHERE client_id = ? AND source = 'quick' ORDER BY created_at DESC",
        (client_id,)
    ).fetchall()
    linked = db.execute(
        "SELECT * FROM client_postiz_groups WHERE client_id = ?", (client_id,)
    ).fetchone()
    db.close()

    channels, channels_error = _get_schedulable_channels(linked)

    return render_template(
        "quick_posts.html",
        clients=clients,
        active_client=active_client,
        quick_posts=quick_posts,
        styles=QUICK_POST_STYLES,
        lengths=QUICK_POST_LENGTHS,
        linked=linked,
        channels=channels,
        channels_error=channels_error,
    )


@app.route("/clients/<int:client_id>/quick-posts/new", methods=["POST"])
def quick_posts_new(client_id):
    """Creates a quick post: point at a Drive folder, describe what's in it,
    Hemingway writes the caption in that client's voice. No Degas transcript
    involved -- Section 9's 'Quick-post design, finalized.'"""
    drive_url = request.form.get("drive_url", "").strip()
    notes = request.form.get("notes", "").strip()
    style = request.form.get("style", "conversational")
    length = request.form.get("length", "short")
    context = request.form.get("context", "").strip()

    if not notes:
        return render_template("error.html", message="Notes can't be empty -- Hemingway needs something to write about."), 400

    try:
        result = hemingway_client.generate_single_post(client_id, notes, style, length, context)
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502

    db = get_db()
    db.execute(
        """INSERT INTO posts (client_id, source, caption, media_ref, status, hemingway_post_id)
           VALUES (?, 'quick', ?, ?, 'draft', ?)""",
        (client_id, result["body"], drive_url, result["post_id"])
    )
    db.commit()
    db.close()
    return redirect(url_for("quick_posts_view", client_id=client_id))


@app.route("/quick-posts/<int:post_id>/regenerate", methods=["POST"])
def quick_posts_regenerate(post_id):
    client_id = request.form.get("client_id", type=int)
    instruction = request.form.get("instruction", "").strip()

    db = get_db()
    post = db.execute("SELECT * FROM posts WHERE id = ?", (post_id,)).fetchone()
    if not post or not post["hemingway_post_id"]:
        db.close()
        return render_template("error.html", message="Can't regenerate -- no linked Hemingway post found."), 400

    try:
        result = hemingway_client.rewrite_post(post["hemingway_post_id"], instruction)
    except hemingway_client.HemingwayError as e:
        db.close()
        return render_template("error.html", message=str(e)), 502

    db.execute("UPDATE posts SET caption = ? WHERE id = ?", (result["body"], post_id))
    db.commit()
    db.close()
    return _post_redirect(post, client_id)


@app.route("/quick-posts/<int:post_id>/advance", methods=["POST"])
def quick_posts_advance(post_id):
    """Minimal 3-state tracking: Draft -> Scheduled -> Published (Section 9).
    Manual for now -- becomes a real Postiz push once task #12 is built."""
    client_id = request.form.get("client_id", type=int)
    db = get_db()
    post = db.execute("SELECT * FROM posts WHERE id = ?", (post_id,)).fetchone()
    target = QUICK_POST_TRANSITIONS.get(post["status"]) if post else None
    if target:
        db.execute("UPDATE posts SET status = ? WHERE id = ?", (target, post_id))
        db.commit()
    db.close()
    return _post_redirect(post, client_id)


@app.route("/quick-posts/<int:post_id>/schedule-to-postiz", methods=["POST"])
def quick_posts_schedule_to_postiz(post_id):
    """Real draft -> scheduled transition (task #13): pushes the post to
    Postiz on the channels you pick, at the time you pick, instead of just
    flipping a local status flag. Postiz becomes the source of truth for
    whether/when this actually goes out.

    Task #27 (Ben's ask 7/24, "I want the copy to connect to its correlating
    video"): project-sourced posts now carry the Degas clip_id they were
    written from. If any selected channel needs media (YouTube/Instagram --
    postiz_client.MEDIA_REQUIRED_IDENTIFIERS), this fetches that exact clip's
    EXPORTED (captioned) video from Degas and uploads it to Postiz once,
    reusing the same upload across every media-required channel in this
    submission. If the clip isn't exported yet, this blocks with a clear
    message rather than silently falling back to the uncaptioned original --
    Ben's explicit call, not a judgment call made here."""
    client_id = request.form.get("client_id", type=int)
    channel_ids = request.form.getlist("channel_ids")
    send_at = request.form.get("send_at", "").strip()
    youtube_title = request.form.get("youtube_title", "").strip()

    if not channel_ids or not send_at:
        return render_template("error.html", message="Pick at least one channel and a send time."), 400

    db = get_db()
    post = db.execute("SELECT * FROM posts WHERE id = ?", (post_id,)).fetchone()
    linked = db.execute("SELECT * FROM client_postiz_groups WHERE client_id = ?", (client_id,)).fetchone()
    proj = None
    if post and post["project_id"]:
        proj = db.execute("SELECT * FROM projects WHERE id = ?", (post["project_id"],)).fetchone()
    if not post or not linked:
        db.close()
        return render_template("error.html", message="Can't schedule -- no linked Postiz group for this client."), 400

    try:
        integrations = postiz_client.list_integrations(linked["postiz_group_id"])
        by_id = {i["id"]: i for i in integrations}

        needs_media = any(
            by_id.get(cid, {}).get("identifier") in postiz_client.MEDIA_REQUIRED_IDENTIFIERS
            for cid in channel_ids
        )
        uploaded_media = None
        if needs_media:
            if not post["clip_id"]:
                db.close()
                return render_template("error.html", message="This post has no linked video -- only project-sourced posts (written from a reviewed clip) can schedule to YouTube/Instagram."), 400
            if not proj or not proj["degas_project_id"]:
                db.close()
                return render_template("error.html", message="This post's project isn't linked to Degas -- can't fetch its video."), 400
            try:
                clip_status = degas_client.get_clip_status(proj["degas_project_id"], post["clip_id"])
            except degas_client.DegasError as e:
                db.close()
                return render_template("error.html", message=str(e)), 502
            if clip_status.get("status") != "exported":
                db.close()
                return render_template("error.html", message="This post's clip isn't exported yet -- export it in Caption Review before scheduling to YouTube/Instagram."), 400
            try:
                degas_resp = degas_client.download_clip(proj["degas_project_id"], post["clip_id"])
                video_bytes = degas_resp.content
            except degas_client.DegasError as e:
                db.close()
                return render_template("error.html", message=str(e)), 502
            uploaded_media = postiz_client.upload_file(
                io.BytesIO(video_bytes), f"clip_{post['clip_id']}.mp4", "video/mp4"
            )

        posts_payload = []
        for cid in channel_ids:
            integ = by_id.get(cid)
            if not integ:
                continue
            image = None
            extra = None
            if integ["identifier"] in postiz_client.MEDIA_REQUIRED_IDENTIFIERS and uploaded_media:
                image = [{"id": uploaded_media["id"], "path": uploaded_media["path"]}]
            if integ["identifier"] == "youtube":
                extra = {"title": youtube_title}
            posts_payload.append(postiz_client.build_post_item(cid, integ["identifier"], post["caption"], image=image, extra=extra))
        if not posts_payload:
            db.close()
            return render_template("error.html", message="None of the selected channels could be matched to a connected integration."), 400

        # send_at comes from an HTML datetime-local input ("YYYY-MM-DDTHH:MM"),
        # which has no timezone. Treated as UTC for now -- mapping it to the
        # client's actual local timezone is a real gap, not handled yet.
        date_iso = f"{send_at}:00.000Z" if len(send_at) == 16 else send_at
        result = postiz_client.create_post("schedule", date_iso, posts_payload)
    except postiz_client.PostizError as e:
        db.close()
        return render_template("error.html", message=str(e)), 502

    postiz_ids = ",".join(str(r.get("postId")) for r in result)
    db.execute(
        "UPDATE posts SET status = 'scheduled', postiz_post_id = ?, scheduled_for = ? WHERE id = ?",
        (postiz_ids, send_at, post_id)
    )
    db.commit()
    db.close()
    return _post_redirect(post, client_id)


@app.route("/quick-posts/<int:post_id>/edit", methods=["POST"])
def quick_posts_edit(post_id):
    """Manual caption edit -- you don't always need Hemingway to touch it,
    sometimes a typo fix is faster by hand."""
    client_id = request.form.get("client_id", type=int)
    caption = request.form.get("caption", "")
    db = get_db()
    post = db.execute("SELECT * FROM posts WHERE id = ?", (post_id,)).fetchone()
    db.execute("UPDATE posts SET caption = ? WHERE id = ?", (caption, post_id))
    db.commit()
    db.close()
    return _post_redirect(post, client_id)


# ── Postiz setup (task #12) ──────────────────────────────────────────────────
@app.route("/clients/<int:client_id>/calendar")
def calendar_view(client_id):
    """Single-client calendar (task #13, Section 6: 'strictly single-client
    view, no cross-client rollup'). Reads Postiz's own /posts endpoint live
    for the selected month -- Postiz's `state` field (QUEUE/PUBLISHED/ERROR/
    DRAFT) is the real status, not something Studio tracks separately.

    Renders a full month grid (Sun-Sat, always 6 weeks) even when empty, and
    every day cell can open the create-post form scoped to that date -- Ben's
    ask (7/24) after seeing Hey Orca's calendar, where "Add Post" lives on the
    calendar itself rather than needing a separate Quick Posts screen first."""
    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502
    active_client = next((c for c in clients if c["id"] == client_id), None)
    if not active_client:
        return redirect(url_for("dashboard"))

    today = datetime.now(timezone.utc)
    month_str = request.args.get("month", "")
    try:
        year, month = map(int, month_str.split("-"))
    except ValueError:
        year, month = today.year, today.month
    month_key = f"{year:04d}-{month:02d}"

    start = datetime(year, month, 1, tzinfo=timezone.utc)
    last_day = cal_module.monthrange(year, month)[1]
    end = datetime(year, month, last_day, 23, 59, 59, tzinfo=timezone.utc)
    prev_month = (start - timedelta(days=1)).strftime("%Y-%m")
    next_month = (end + timedelta(days=1)).strftime("%Y-%m")

    # Sun-Sat grid, always full weeks (padded with adjacent-month dates),
    # same shape as Hey Orca's month view -- firstweekday=6 makes Sunday the
    # first column.
    weeks = cal_module.Calendar(firstweekday=6).monthdatescalendar(year, month)

    db = get_db()
    linked = db.execute(
        "SELECT * FROM client_postiz_groups WHERE client_id = ?", (client_id,)
    ).fetchone()
    db.close()

    # Media-capable set here (not the text-only default) -- the calendar's
    # Add Post modal has a file upload step, so Instagram/YouTube are real
    # options here even though Quick Posts doesn't offer them.
    channels, channels_error = _get_schedulable_channels(linked, postiz_client.MEDIA_CAPABLE_IDENTIFIERS)

    days = {}
    postiz_error = None
    if not linked:
        postiz_error = f"{active_client['name']} isn't linked to a Postiz group yet -- set that up first."
    else:
        try:
            resp = postiz_client.list_posts(
                start.strftime("%Y-%m-%dT00:00:00.000Z"),
                end.strftime("%Y-%m-%dT23:59:59.000Z"),
                customer=linked["postiz_group_id"],
            )
            for p in resp.get("posts", []):
                day = p["publishDate"][:10]
                time_str = p["publishDate"][11:16]
                days.setdefault(day, []).append({
                    "id": p["id"],
                    "content": p["content"],
                    "time": time_str,
                    "state": p["state"],
                    "channel_name": p.get("integration", {}).get("name", ""),
                    "channel_platform": p.get("integration", {}).get("providerIdentifier", ""),
                    "release_url": p.get("releaseURL"),
                })
        except postiz_client.PostizError as e:
            postiz_error = str(e)

    return render_template(
        "calendar.html",
        clients=clients,
        active_client=active_client,
        linked=linked,
        channels=channels,
        channels_error=channels_error,
        media_required_identifiers=list(postiz_client.MEDIA_REQUIRED_IDENTIFIERS),
        days=days,
        weeks=weeks,
        current_month=month,
        today_str=today.strftime("%Y-%m-%d"),
        postiz_error=postiz_error,
        month_label=start.strftime("%B %Y"),
        month_key=month_key,
        prev_month=prev_month,
        next_month=next_month,
    )


@app.route("/clients/<int:client_id>/calendar/create-post", methods=["POST"])
def calendar_create_post(client_id):
    """Direct "Add Post" from the calendar (task #13 follow-up, 7/24): types
    a caption and picks channels right on the calendar, like Hey Orca's
    'Create Post(s)' modal, instead of going through Quick Posts first. No
    Hemingway involved here -- this is for a caption you're writing yourself;
    Quick Posts is still the path when you want Hemingway's help drafting it.

    Now handles an optional media file (task #18): uploaded once to Postiz
    and attached to every selected channel that accepts it, since Instagram
    and YouTube both require real media -- see
    postiz_client.MEDIA_REQUIRED_IDENTIFIERS."""
    caption = request.form.get("caption", "").strip()
    channel_ids = request.form.getlist("channel_ids")
    send_at = request.form.get("send_at", "").strip()
    month_key = request.form.get("month_key", "")
    youtube_title = request.form.get("youtube_title", "").strip()
    # getlist, not get -- the file input now allows multiple selections
    # (Ben's ask, 2026-08-24: multi-image carousel posts). Postiz's `image`
    # field on a post was always a list; Studio just never sent more than one
    # entry before. Order matters -- it's carousel order and cover order --
    # so files are uploaded and appended in the order the browser sent them.
    media_files = [f for f in request.files.getlist("media") if f and f.filename]

    if not caption or not channel_ids or not send_at:
        return render_template("error.html", message="A caption, at least one channel, and a send time are all required."), 400

    db = get_db()
    linked = db.execute("SELECT * FROM client_postiz_groups WHERE client_id = ?", (client_id,)).fetchone()
    if not linked:
        db.close()
        return render_template("error.html", message="This client isn't linked to a Postiz group yet -- set that up first."), 400

    try:
        integrations = postiz_client.list_integrations(linked["postiz_group_id"])
        by_id = {i["id"]: i for i in integrations}
        selected = [by_id[cid] for cid in channel_ids if cid in by_id]
        if not selected:
            db.close()
            return render_template("error.html", message="None of the selected channels could be matched to a connected integration."), 400

        needs_media = any(i["identifier"] in postiz_client.MEDIA_REQUIRED_IDENTIFIERS for i in selected)
        needs_youtube = any(i["identifier"] == "youtube" for i in selected)
        if needs_youtube and len(media_files) > 1:
            db.close()
            return render_template("error.html", message="YouTube accepts exactly one video file, not multiple -- remove the extra files or unselect YouTube."), 400

        image = []
        if media_files:
            for f in media_files:
                uploaded = postiz_client.upload_file(f.stream, f.filename, f.content_type)
                image.append({"id": uploaded["id"], "path": uploaded["path"]})
        elif needs_media:
            db.close()
            needed = [i["identifier"] for i in selected if i["identifier"] in postiz_client.MEDIA_REQUIRED_IDENTIFIERS]
            return render_template("error.html", message=f"{', '.join(needed)} requires an image or video attached -- none was uploaded."), 400

        posts_payload = [
            postiz_client.build_post_item(
                i["id"], i["identifier"], caption,
                image=image, extra={"title": youtube_title},
            )
            for i in selected
        ]

        date_iso = f"{send_at}:00.000Z" if len(send_at) == 16 else send_at
        result = postiz_client.create_post("schedule", date_iso, posts_payload)
    except postiz_client.PostizError as e:
        db.close()
        return render_template("error.html", message=str(e)), 502

    postiz_ids = ",".join(str(r.get("postId")) for r in result)
    db.execute(
        """INSERT INTO posts (client_id, source, caption, status, postiz_post_id, scheduled_for)
           VALUES (?, 'quick', ?, 'scheduled', ?, ?)""",
        (client_id, caption, postiz_ids, send_at)
    )
    db.commit()
    db.close()
    return redirect(url_for("calendar_view", client_id=client_id, month=month_key))


@app.route("/clients/<int:client_id>/calendar/posts/<post_id>/delete", methods=["POST"])
def calendar_delete_post(client_id, post_id):
    """Delete a scheduled/draft/error post from the calendar (Ben's ask,
    2026-08-24: no way to edit or move posts once scheduled). Postiz's public
    API has no endpoint to change a post's time or content -- Update Post
    Settings only merges provider settings and explicitly leaves content and
    publishDate untouched -- so delete is the only safe, real "undo" Studio
    can offer. For editing or moving a post, use Postiz's own calendar
    directly (drag-and-drop reschedule works there). See postiz_client.delete_post."""
    month_key = request.form.get("month_key", "")
    try:
        postiz_client.delete_post(post_id)
    except postiz_client.PostizError as e:
        return render_template("error.html", message=str(e)), 502
    return redirect(url_for("calendar_view", client_id=client_id, month=month_key))


@app.route("/clients/<int:client_id>/postiz-setup")
def postiz_setup_view(client_id):
    """Links a Hemingway client to a Postiz customer group. Real scheduling
    (picking a channel + send time and pushing a post) lands with the
    calendar (task #13) -- this screen is just the one-time account linkage
    each client needs before that can work."""
    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502
    active_client = next((c for c in clients if c["id"] == client_id), None)
    if not active_client:
        return redirect(url_for("dashboard"))

    connected = None
    groups = []
    postiz_error = None
    try:
        connected = postiz_client.is_connected()
        if connected:
            groups = postiz_client.list_groups()
    except postiz_client.PostizError as e:
        postiz_error = str(e)

    db = get_db()
    linked = db.execute(
        "SELECT * FROM client_postiz_groups WHERE client_id = ?", (client_id,)
    ).fetchone()
    db.close()

    integrations = []
    integrations_error = None
    if linked:
        try:
            integrations = postiz_client.list_integrations(linked["postiz_group_id"])
        except postiz_client.PostizError as e:
            integrations_error = str(e)

    return render_template(
        "postiz_setup.html",
        clients=clients,
        active_client=active_client,
        connected=connected,
        postiz_error=postiz_error,
        groups=groups,
        linked=linked,
        integrations=integrations,
        integrations_error=integrations_error,
    )


@app.route("/clients/<int:client_id>/postiz-setup/link", methods=["POST"])
def postiz_setup_link(client_id):
    group_id = request.form.get("group_id", "").strip()
    group_name = request.form.get("group_name", "").strip()
    if not group_id:
        return redirect(url_for("postiz_setup_view", client_id=client_id))

    db = get_db()
    db.execute(
        """INSERT INTO client_postiz_groups (client_id, postiz_group_id, postiz_group_name, updated_at)
           VALUES (?, ?, ?, CURRENT_TIMESTAMP)
           ON CONFLICT(client_id) DO UPDATE SET
             postiz_group_id = excluded.postiz_group_id,
             postiz_group_name = excluded.postiz_group_name,
             updated_at = CURRENT_TIMESTAMP""",
        (client_id, group_id, group_name)
    )
    db.commit()
    db.close()
    return redirect(url_for("postiz_setup_view", client_id=client_id))


@app.route("/clients/<int:client_id>/nextcloud-setup")
def nextcloud_setup_view(client_id):
    """One-time linkage: which top-level Cloud KMG folder belongs to this
    client (task #28, Ben's ask 8/3). Saving here also creates
    /<folder>/incoming and /<folder>/captioned if they don't exist yet, so
    there's nothing else to set up before the project page's Cloud KMG card
    and Archive-to-Cloud buttons work."""
    try:
        clients = hemingway_client.get_clients()
    except hemingway_client.HemingwayError as e:
        return render_template("error.html", message=str(e)), 502
    active_client = next((c for c in clients if c["id"] == client_id), None)
    if not active_client:
        return redirect(url_for("dashboard"))

    linked = _get_nextcloud_folder(client_id)

    incoming_files = []
    imported_files = []
    captioned_files = []
    nextcloud_error = None
    if linked:
        try:
            incoming_files = nextcloud_client.list_folder(f"{linked['folder_name']}/incoming")
            imported_files = nextcloud_client.list_folder(f"{linked['folder_name']}/imported")
            captioned_files = nextcloud_client.list_folder(f"{linked['folder_name']}/captioned")
        except nextcloud_client.NextcloudError as e:
            nextcloud_error = str(e)

    return render_template(
        "nextcloud_setup.html",
        clients=clients,
        active_client=active_client,
        linked=linked,
        incoming_files=incoming_files,
        imported_files=imported_files,
        captioned_files=captioned_files,
        nextcloud_error=nextcloud_error,
    )


@app.route("/clients/<int:client_id>/nextcloud-setup/link", methods=["POST"])
def nextcloud_setup_link(client_id):
    folder_name = request.form.get("folder_name", "").strip().strip("/")
    if not folder_name:
        return redirect(url_for("nextcloud_setup_view", client_id=client_id))

    try:
        nextcloud_client.ensure_folder(f"{folder_name}/incoming")
        nextcloud_client.ensure_folder(f"{folder_name}/captioned")
        nextcloud_client.ensure_folder(f"{folder_name}/imported")
    except nextcloud_client.NextcloudError as e:
        return render_template("error.html", message=str(e)), 502

    db = get_db()
    db.execute(
        """INSERT INTO client_nextcloud_folders (client_id, folder_name, updated_at)
           VALUES (?, ?, CURRENT_TIMESTAMP)
           ON CONFLICT(client_id) DO UPDATE SET
             folder_name = excluded.folder_name,
             updated_at = CURRENT_TIMESTAMP""",
        (client_id, folder_name)
    )
    db.commit()
    db.close()
    return redirect(url_for("nextcloud_setup_view", client_id=client_id))


# ── Startup ───────────────────────────────────────────────────────────────────
_initialized = False


@app.before_request
def ensure_initialized():
    global _initialized
    if not _initialized:
        init_db()
        _initialized = True


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
